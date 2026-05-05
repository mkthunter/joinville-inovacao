"""Gera o DOCX do Decreto dos APIs (não publicado na plataforma)."""
import sys
sys.path.insert(0, 'scripts')

from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.enum.text

# Importar funções do script principal
spec_path = Path('scripts/gerar-decretos-word.py')
import importlib.util
spec = importlib.util.spec_from_file_location("gen_word", spec_path)
gw = importlib.util.module_from_spec(spec)
spec.loader.exec_module(gw)

ROOT = Path(__file__).parent.parent
OUT_DIR = ROOT / "public" / "pdfs"
OUT_DIR.mkdir(parents=True, exist_ok=True)

doc = gw.create_base_doc()

# Capa
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run("PREFEITURA MUNICIPAL DE JOINVILLE")
run.bold = True
run.font.size = Pt(14)
gw.set_paragraph_format(p_title, space_before=0, space_after=6)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run("MINUTA DE DECRETO")
r2.font.size = Pt(12)
gw.set_paragraph_format(p_sub, space_before=0, space_after=4)

p_nota = doc.add_paragraph()
p_nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_nota.add_run("Versão consolidada para análise jurídica")
r3.italic = True
r3.font.size = Pt(10)
gw.set_paragraph_format(p_nota, space_before=0, space_after=4)

p_org = doc.add_paragraph()
p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p_org.add_run("BRZ Capacitação × Consultoria SEBRAE/SC · maio/2026")
r4.font.size = Pt(10)
gw.set_paragraph_format(p_org, space_before=0, space_after=20)

# Título
p_titulo_dec = doc.add_paragraph()
p_titulo_dec.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = p_titulo_dec.add_run("DECRETO DOS ARRANJOS PROMOTORES DE INOVAÇÃO (APIs)")
rt.bold = True
rt.font.size = Pt(12)
gw.set_paragraph_format(p_titulo_dec, space_before=20, space_after=18)

apis_items = gw.extract_apis_sections()
print(f"APIs: {len(apis_items)} elementos extraídos")
gw.render_items(doc, apis_items)

out_path = OUT_DIR / "Joinville - Decreto APIs.docx"
doc.save(str(out_path))
print(f"\n✓ Salvo: {out_path}")
