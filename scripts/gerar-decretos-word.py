"""
Gera os 2 documentos Word unificados limpos:
  - decreto-unificado-rota-a.docx  (PII Rota A + APIs)
  - decreto-unificado-rota-b.docx  (PII Rota B + APIs)

Extrai diretamente do index.html (fonte autoritativa).
Exclui: dec-nota-redacao, dec-section__subtitle, dec-section__num,
        dec-caminho__cabecalho, dec-toggle-inline[data-when!=rota].
"""

import re
import sys
from pathlib import Path
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────────
ROOT = Path(__file__).parent.parent
HTML_PATH = ROOT / "index.html"
OUT_DIR = ROOT / "public" / "pdfs"
OUT_DIR.mkdir(parents=True, exist_ok=True)

# ── Parse HTML ─────────────────────────────────────────────────────────────────
with open(HTML_PATH, encoding="utf-8") as f:
    soup = BeautifulSoup(f.read(), "html.parser")


# ── Helpers ────────────────────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    """Remove espaços extras mas preserva conteúdo íntegro."""
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def get_text_for_rota(el: Tag, rota: str) -> str:
    """
    Extrai texto de um elemento, respeitando data-when:
      - spans com data-when == rota: incluir
      - spans com data-when != rota: excluir
      - sem data-when: incluir normalmente
    Exclui dec-nota-redacao.
    """
    parts = []
    for node in el.descendants:
        # Pula notas consultivas
        if isinstance(node, Tag):
            if "dec-nota-redacao" in node.get("class", []):
                continue
            # Tag com data-when errado: pula
            if node.name == "span" and node.has_attr("data-when"):
                if node["data-when"] != rota:
                    continue
        if isinstance(node, NavigableString):
            # Verifica se algum ancestral tem data-when errado
            skip = False
            for parent in node.parents:
                if isinstance(parent, Tag):
                    if "dec-nota-redacao" in parent.get("class", []):
                        skip = True
                        break
                    if parent.name == "span" and parent.has_attr("data-when"):
                        if parent["data-when"] != rota:
                            skip = True
                            break
            if not skip:
                parts.append(str(node))
    return clean_text("".join(parts))


def set_paragraph_format(para, space_before=0, space_after=6, line_spacing=None):
    pPr = para._p.get_or_add_pPr()
    pPr_spacing = pPr.find(qn("w:spacing"))
    if pPr_spacing is None:
        pPr_spacing = OxmlElement("w:spacing")
        pPr.append(pPr_spacing)
    pPr_spacing.set(qn("w:before"), str(space_before * 20))
    pPr_spacing.set(qn("w:after"), str(space_after * 20))
    if line_spacing:
        pPr_spacing.set(qn("w:line"), str(int(line_spacing * 240)))
        pPr_spacing.set(qn("w:lineRule"), "auto")


def add_page_break(doc: Document):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)
    return para


# ── Criar documento Word ───────────────────────────────────────────────────────

def create_base_doc() -> Document:
    doc = Document()

    # Margens (2cm todos os lados)
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Estilo Normal base
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    return doc


def add_decree_title(doc: Document, text: str):
    """DECRETO Nº ___, DE ___ DE ___ DE 2026."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    set_paragraph_format(p, space_before=0, space_after=10)
    return p


def add_ementa(doc: Document, text: str):
    """Ementa em itálico centralizada."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    set_paragraph_format(p, space_before=4, space_after=10)
    return p


def add_chapter_title(doc: Document, text: str):
    """CAPÍTULO I / CAPÍTULO II etc."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    set_paragraph_format(p, space_before=16, space_after=4)
    return p


def add_section_title(doc: Document, text: str):
    """Seção I — Do objeto"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    set_paragraph_format(p, space_before=8, space_after=4)
    return p


def add_body_paragraph(doc: Document, text: str, indent_first: bool = True):
    """Parágrafo de corpo com recuo de 1ª linha."""
    if not text:
        return None
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.size = Pt(12)
    set_paragraph_format(p, space_before=0, space_after=6, line_spacing=1.15)
    return p


def add_separator(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("* * *")
    run.font.size = Pt(10)
    set_paragraph_format(p, space_before=20, space_after=20)
    return p


def add_section_header_divider(doc: Document, cap_num: str, cap_title: str):
    """Cabeçalho de capítulo: CAPÍTULO I\nDAS DISPOSIÇÕES GERAIS"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p.add_run(f"{cap_num}\n")
    run1.bold = True
    run1.font.size = Pt(12)
    run2 = p.add_run(cap_title.upper())
    run2.bold = True
    run2.font.size = Pt(12)
    set_paragraph_format(p, space_before=18, space_after=6)
    return p


# ── Extrair seções do decreto ──────────────────────────────────────────────────

def extract_pii_sections(rota: str) -> list:
    """
    Retorna lista de (tipo, texto) extraídos da seção caminho-{rota}-completo.
    tipos: 'ementa', 'autor', 'bloco_title', 'artigo', 'paragrafo',
           'inciso_item', 'section_title', 'pmj_item', 'considerando_item'
    """
    import copy

    anchor = f"caminho-{rota}-completo"
    section = copy.copy(soup.find(id=anchor))
    if not section:
        print(f"  ERRO: #{anchor} não encontrado", file=sys.stderr)
        return []

    # Remove blocos dec-caminho que NÃO são do rota atual
    # Identifica pela classe modificadora dec-caminho--a / dec-caminho--b
    rota_class = f"dec-caminho--{rota}"
    for bloco in section.find_all(class_="dec-caminho"):
        classes = bloco.get("class", [])
        if rota_class not in classes:
            bloco.decompose()

    items = []

    for el in section.find_all(True, recursive=True):
        classes = el.get("class", [])
        tag = el.name

        # Pular notas consultivas
        if "dec-nota-redacao" in classes:
            continue
        # Pular cabeçalhos de caminho (A/B labels)
        if "dec-caminho__cabecalho" in classes:
            continue
        # Pular subtítulos de seção (metadado editorial)
        if "dec-section__subtitle" in classes:
            continue
        if "dec-section__num" in classes:
            continue

        # Ementa
        if "dec-norma__ementa" in classes:
            txt = get_text_for_rota(el, rota)
            if txt:
                tipo = "ementa_italico" if "dec-norma__ementa--italico" in classes else "ementa"
                items.append((tipo, txt))
            continue

        # Autor
        if "dec-norma__autor" in classes:
            txt = get_text_for_rota(el, rota)
            if txt:
                items.append(("autor", txt))
            continue

        # Título de bloco/seção (CONSIDERANDO, DECRETA, Seção I etc.)
        if "dec-norma__bloco-title" in classes:
            txt = clean_text(el.get_text())
            if txt:
                items.append(("bloco_title", txt))
            continue

        # Título de seção (capítulo)
        if "dec-section__title" in classes:
            txt = clean_text(el.get_text())
            if txt:
                items.append(("section_title", txt))
            continue

        # Artigo
        if "dec-artigo" in classes:
            txt = get_text_for_rota(el, rota)
            if txt:
                items.append(("artigo", txt))
            continue

        # Parágrafo
        if "dec-paragrafo" in classes:
            txt = get_text_for_rota(el, rota)
            if txt:
                items.append(("paragrafo", txt))
            continue

        # Considerandos (ol.dec-considerandos)
        if "dec-considerandos" in classes:
            for li in el.find_all("li", recursive=False):
                txt = get_text_for_rota(li, rota)
                # Filtra placeholder "(sem considerando específico...)"
                if txt and not txt.startswith("(sem considerando"):
                    items.append(("considerando", txt))
            continue

        # Incisos (lista)
        if "dec-incisos" in classes:
            for li in el.find_all("li"):
                txt = get_text_for_rota(li, rota)
                if txt:
                    items.append(("inciso", txt))
            continue

        # Lista PMJ
        if "dec-pmj-lista" in classes:
            for item in el.find_all(class_="dec-pmj-item"):
                num_el = item.find(class_="dec-pmj-item__num")
                num = clean_text(num_el.get_text()) if num_el else ""
                # texto é tudo menos o num
                if num_el:
                    num_el.decompose()
                txt = get_text_for_rota(item, rota)
                if txt:
                    items.append(("pmj_item", f"{num} {txt}".strip()))
            continue

    return items


def extract_apis_sections() -> list:
    """Extrai o Decreto dos APIs (não tem variante A/B)."""
    import copy
    page = copy.copy(soup.find(id="page-decreto-apis"))
    if not page:
        print("  ERRO: #page-decreto-apis não encontrado", file=sys.stderr)
        return []
    # Remove elementos de navegação antes de extrair
    for unwanted in page.find_all(class_=["decreto-downloads", "decreto-overview",
                                           "decreto-toc-direito", "page-header",
                                           "decreto-apis-nav"]):
        unwanted.decompose()
    for unwanted in page.find_all(class_=lambda c: c and any(
        x in c for x in ["decreto-downloads", "decreto-overview", "dec-toc", "nav"]
    )):
        unwanted.decompose()

    items = []

    for el in page.find_all(True, recursive=True):
        classes = el.get("class", [])

        if "dec-nota-redacao" in classes:
            continue
        if "dec-section__subtitle" in classes:
            continue
        if "dec-section__num" in classes:
            continue
        if "dec-caminho__cabecalho" in classes:
            continue
        # Excluir nav interna da página e seções de download
        if "decreto-downloads" in classes or "decreto-overview" in classes:
            continue
        if "decreto-toc-direito" in classes or "decreto-overview__atalho" in classes:
            continue

        if "dec-norma__ementa" in classes:
            txt = clean_text(el.get_text())
            if txt:
                tipo = "ementa_italico" if "dec-norma__ementa--italico" in classes else "ementa"
                items.append((tipo, txt))
            continue

        if "dec-norma__autor" in classes:
            txt = clean_text(el.get_text())
            if txt:
                items.append(("autor", txt))
            continue

        if "dec-norma__bloco-title" in classes:
            txt = clean_text(el.get_text())
            if txt:
                items.append(("bloco_title", txt))
            continue

        if "dec-section__title" in classes:
            txt = clean_text(el.get_text())
            if txt:
                items.append(("section_title", txt))
            continue

        if "dec-artigo" in classes:
            txt = clean_text(el.get_text())
            if txt:
                items.append(("artigo", txt))
            continue

        if "dec-paragrafo" in classes:
            txt = clean_text(el.get_text())
            if txt:
                items.append(("paragrafo", txt))
            continue

        if "dec-considerandos" in classes:
            for li in el.find_all("li", recursive=False):
                txt = clean_text(li.get_text())
                if txt and not txt.startswith("(sem considerando"):
                    items.append(("considerando", txt))
            continue

        if "dec-incisos" in classes:
            for li in el.find_all("li"):
                txt = clean_text(li.get_text())
                if txt:
                    items.append(("inciso", txt))
            continue

        if "dec-pmj-lista" in classes:
            for item in el.find_all(class_="dec-pmj-item"):
                num_el = item.find(class_="dec-pmj-item__num")
                num = clean_text(num_el.get_text()) if num_el else ""
                if num_el:
                    num_el.decompose()
                txt = clean_text(item.get_text())
                if txt:
                    items.append(("pmj_item", f"{num} {txt}".strip()))
            continue

    return items


# ── Renderizar items no documento Word ────────────────────────────────────────

def render_items(doc: Document, items: list):
    seen = set()
    prev_tipo = None

    for tipo, texto in items:
        # Deduplicar: mesmo tipo+texto consecutivo = skip
        key = (tipo, texto)
        if key in seen and tipo in ("ementa", "ementa_italico", "autor", "section_title"):
            continue
        seen.add(key)

        if tipo == "ementa":
            add_decree_title(doc, texto)

        elif tipo == "ementa_italico":
            add_ementa(doc, texto)

        elif tipo == "autor":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(texto)
            run.font.size = Pt(12)
            set_paragraph_format(p, space_before=10, space_after=6)

        elif tipo == "section_title":
            add_chapter_title(doc, texto)

        elif tipo == "bloco_title":
            add_section_title(doc, texto)

        elif tipo == "artigo":
            add_body_paragraph(doc, texto, indent_first=True)

        elif tipo == "paragrafo":
            add_body_paragraph(doc, texto, indent_first=True)

        elif tipo == "considerando":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.25)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            run = p.add_run(texto)
            run.font.size = Pt(12)
            set_paragraph_format(p, space_before=0, space_after=5, line_spacing=1.15)

        elif tipo == "inciso":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            run = p.add_run(texto)
            run.font.size = Pt(12)
            set_paragraph_format(p, space_before=0, space_after=3, line_spacing=1.15)

        elif tipo == "pmj_item":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            run = p.add_run(texto)
            run.font.size = Pt(11)
            set_paragraph_format(p, space_before=0, space_after=4)

        prev_tipo = tipo


# ── Montar documento por rota ─────────────────────────────────────────────────

import docx.enum.text

def gerar_documento(rota: str):
    label = rota.upper()
    print(f"\n{'='*60}")
    print(f"  Gerando Rota {label}...")
    print(f"{'='*60}")

    doc = create_base_doc()

    # ── Capa ──────────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("PREFEITURA MUNICIPAL DE JOINVILLE")
    run.bold = True
    run.font.size = Pt(14)
    set_paragraph_format(p_title, space_before=0, space_after=6)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p_sub.add_run("MINUTA DE DECRETO")
    run2.font.size = Pt(12)
    set_paragraph_format(p_sub, space_before=0, space_after=4)

    p_nota = doc.add_paragraph()
    p_nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p_nota.add_run("Versão consolidada para análise jurídica")
    run3.font.size = Pt(10)
    run3.italic = True
    set_paragraph_format(p_nota, space_before=0, space_after=4)

    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p_org.add_run("BRZ Capacitação × Consultoria SEBRAE/SC · abril/2026")
    run4.font.size = Pt(10)
    set_paragraph_format(p_org, space_before=0, space_after=40)

    # ── Título da minuta ───────────────────────────────────────────────────
    p_titulo_dec = doc.add_paragraph()
    p_titulo_dec.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_td = p_titulo_dec.add_run("DECRETO DO PROGRAMA MUNICIPAL DE INCENTIVO À INOVAÇÃO")
    run_td.bold = True
    run_td.font.size = Pt(12)
    set_paragraph_format(p_titulo_dec, space_before=20, space_after=18)

    pii_items = extract_pii_sections(rota)
    print(f"  PII Rota {label}: {len(pii_items)} elementos extraídos")
    render_items(doc, pii_items)

    # ── Salvar ─────────────────────────────────────────────────────────────
    nome = f"Joinville - R-{label} - Decreto PII.docx"
    out_path = OUT_DIR / nome
    doc.save(str(out_path))
    print(f"\n  ✓ Salvo: {out_path}")
    return out_path


# ── Main ───────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    gerar_documento("a")
    gerar_documento("b")
    print("\nConcluído.")
