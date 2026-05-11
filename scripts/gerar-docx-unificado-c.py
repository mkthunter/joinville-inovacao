"""Gera DOCX unificado do Caminho C: Decreto PII + Decreto APIs em sequência.

Layout oficial:
  - Capa única (Prefeitura, BRZ × SEBRAE, data)
  - Decreto do Programa (PII-C) integral
  - Quebra de página
  - Decreto dos APIs por nomeação (API-C) integral

Times New Roman 12pt, margens 2.5/3.0/2.5/2.5 cm, justificado.
"""
import sys
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import importlib.util

# Importar funções do script principal (gerar-decretos-word.py)
ROOT = Path(__file__).parent.parent
spec_path = Path(__file__).parent / "gerar-decretos-word.py"
spec = importlib.util.spec_from_file_location("gen_word", spec_path)
gw = importlib.util.module_from_spec(spec)
spec.loader.exec_module(gw)

HTML_PATH = ROOT / "index.html"
OUT_DIR = ROOT / "public" / "pdfs"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def clean_text(text):
    import re
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def extract_apis_c_sections():
    """Extrai exclusivamente o bloco #caminho-apis-c-completo."""
    import copy
    with open(HTML_PATH, encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "html.parser")

    block = copy.copy(soup.find(id="caminho-apis-c-completo"))
    if not block:
        print("  ERRO: #caminho-apis-c-completo não encontrado", file=sys.stderr)
        return []

    # Remove elementos não sancionáveis
    for unwanted_class in ["dec-nota-redacao", "resumo-executivo",
                           "caminho-completo__header", "caminho-completo__aviso-orientacoes"]:
        for el in block.find_all(class_=unwanted_class):
            el.decompose()

    items = []
    for el in block.find_all(True, recursive=True):
        classes = el.get("class", [])

        if "dec-nota-redacao" in classes:
            continue
        if "dec-section__subtitle" in classes:
            continue
        if "dec-section__num" in classes:
            continue

        if "dec-norma__ementa" in classes:
            txt = clean_text(el.get_text())
            if txt:
                tipo = "ementa_italico" if "dec-norma__ementa--italico" in classes else "ementa"
                items.append((tipo, txt))
            continue

        if "dec-norma__autor" in classes:
            txt = clean_text(el.get_text())
            if txt:
                items.append(("autor", txt))
            continue

        if "dec-norma__bloco-title" in classes:
            txt = clean_text(el.get_text())
            if txt:
                items.append(("bloco_title", txt))
            continue

        if "dec-section__title" in classes:
            txt = clean_text(el.get_text())
            if txt:
                items.append(("section_title", txt))
            continue

        if "dec-artigo" in classes:
            txt = clean_text(el.get_text())
            if txt:
                items.append(("artigo", txt))
            continue

        if "dec-paragrafo" in classes:
            txt = clean_text(el.get_text())
            if txt:
                items.append(("paragrafo", txt))
            continue

        if "dec-considerandos" in classes:
            for li in el.find_all("li", recursive=False):
                txt = clean_text(li.get_text())
                if txt and not txt.startswith("(sem considerando"):
                    items.append(("considerando", txt))
            continue

        if "dec-incisos" in classes:
            for li in el.find_all("li"):
                txt = clean_text(li.get_text())
                if txt:
                    items.append(("inciso", txt))
            continue

    return items


def add_page_break(doc):
    """Quebra de página real."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    return p


def add_separator_title(doc, text):
    """Título de separação entre Decretos (centralizado, bold, maiúsculas)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(13)
    gw.set_paragraph_format(p, space_before=0, space_after=18)
    return p


# ── Montar documento unificado ────────────────────────────────────────────────

doc = gw.create_base_doc()

# ───── CAPA ÚNICA ─────
p_org_top = doc.add_paragraph()
p_org_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_org_top.add_run("PREFEITURA MUNICIPAL DE JOINVILLE")
run.bold = True
run.font.size = Pt(14)
gw.set_paragraph_format(p_org_top, space_before=0, space_after=8)

p_subtitle = doc.add_paragraph()
p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_subtitle.add_run("MINUTA DE DECRETOS")
r2.bold = True
r2.font.size = Pt(12)
gw.set_paragraph_format(p_subtitle, space_before=0, space_after=4)

p_minuta = doc.add_paragraph()
p_minuta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_minuta.add_run("Programa Municipal de Incentivo à Inovação de Joinville (PII/Jlle)")
r3.font.size = Pt(12)
gw.set_paragraph_format(p_minuta, space_before=0, space_after=4)

p_caminho = doc.add_paragraph()
p_caminho.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p_caminho.add_run("Caminho C — Apoio direto pela dotação orçamentária da Secretaria gestora")
r4.italic = True
r4.font.size = Pt(11)
gw.set_paragraph_format(p_caminho, space_before=0, space_after=24)

p_doc_label = doc.add_paragraph()
p_doc_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p_doc_label.add_run("Documento contendo:")
r5.font.size = Pt(11)
gw.set_paragraph_format(p_doc_label, space_before=0, space_after=4)

p_item1 = doc.add_paragraph()
p_item1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r6 = p_item1.add_run("I — Decreto do Programa Municipal de Incentivo à Inovação")
r6.font.size = Pt(11)
gw.set_paragraph_format(p_item1, space_before=0, space_after=2)

p_item2 = doc.add_paragraph()
p_item2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r7 = p_item2.add_run("II — Decreto dos Arranjos Promotores de Inovação (APIs) — Nomeação")
r7.font.size = Pt(11)
gw.set_paragraph_format(p_item2, space_before=0, space_after=24)

p_versao = doc.add_paragraph()
p_versao.alignment = WD_ALIGN_PARAGRAPH.CENTER
r8 = p_versao.add_run("Versão consolidada para análise jurídica")
r8.italic = True
r8.font.size = Pt(10)
gw.set_paragraph_format(p_versao, space_before=0, space_after=4)

p_org_bottom = doc.add_paragraph()
p_org_bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
r9 = p_org_bottom.add_run("BRZ Capacitação × Consultoria SEBRAE/SC · maio/2026")
r9.font.size = Pt(10)
gw.set_paragraph_format(p_org_bottom, space_before=0, space_after=20)

# ───── QUEBRA DE PÁGINA ─────
add_page_break(doc)

# ───── PARTE I — DECRETO PII-C ─────
add_separator_title(doc, "PARTE I — DECRETO DO PROGRAMA MUNICIPAL DE INCENTIVO À INOVAÇÃO")

p_titulo_dec1 = doc.add_paragraph()
p_titulo_dec1.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt1 = p_titulo_dec1.add_run("DECRETO DO PROGRAMA MUNICIPAL DE INCENTIVO À INOVAÇÃO")
rt1.bold = True
rt1.font.size = Pt(12)
gw.set_paragraph_format(p_titulo_dec1, space_before=4, space_after=18)

pii_items = gw.extract_pii_sections("c")
print(f"PII Caminho C: {len(pii_items)} elementos")
gw.render_items(doc, pii_items)

# ───── QUEBRA DE PÁGINA ─────
add_page_break(doc)

# ───── PARTE II — DECRETO API-C ─────
add_separator_title(doc, "PARTE II — DECRETO DOS ARRANJOS PROMOTORES DE INOVAÇÃO (APIs)")

p_titulo_dec2 = doc.add_paragraph()
p_titulo_dec2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt2 = p_titulo_dec2.add_run("DECRETO DOS ARRANJOS PROMOTORES DE INOVAÇÃO (APIs) — POR NOMEAÇÃO")
rt2.bold = True
rt2.font.size = Pt(12)
gw.set_paragraph_format(p_titulo_dec2, space_before=4, space_after=18)

apis_c_items = extract_apis_c_sections()
print(f"APIs Caminho C: {len(apis_c_items)} elementos")
gw.render_items(doc, apis_c_items)

# ───── SALVAR ─────
out_path = OUT_DIR / "Joinville - Caminho C - Decretos PII + APIs.docx"
doc.save(str(out_path))
print(f"\n✓ Salvo: {out_path}")

# Estatísticas
import re
text = "\n".join(p.text for p in doc.paragraphs)
artigos_pii = len([m for m in re.findall(r"\bArt\.?\s+(\d+)º?\b", text) if int(m) <= 65])
artigos_total_unique = len(set(re.findall(r"\bArt\.?\s+(\d+)º?\b", text)))
print(f"  Parágrafos: {len(doc.paragraphs)}")
print(f"  Caracteres: {len(text)}")
print(f"  Artigos identificados (totais): {artigos_total_unique}")
