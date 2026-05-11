"""Gera o DOCX do Decreto dos APIs por nomeação direta — Caminho C.

Extrai exclusivamente do bloco #caminho-apis-c-completo no index.html,
ignorando o decreto APIs A/B (regime ordinário) que vive na mesma página.
"""
import sys
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.enum.text

# Importar funções do script principal
sys.path.insert(0, str(Path(__file__).parent))
spec_path = Path(__file__).parent / "gerar-decretos-word.py"
import importlib.util
spec = importlib.util.spec_from_file_location("gen_word", spec_path)
gw = importlib.util.module_from_spec(spec)
spec.loader.exec_module(gw)

ROOT = Path(__file__).parent.parent
HTML_PATH = ROOT / "index.html"
OUT_DIR = ROOT / "public" / "pdfs"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def clean_text(text):
    import re
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def extract_apis_c_sections():
    """Extrai exclusivamente o bloco #caminho-apis-c-completo (API-C nomeação)."""
    import copy
    with open(HTML_PATH, encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "html.parser")

    block = copy.copy(soup.find(id="caminho-apis-c-completo"))
    if not block:
        print("  ERRO: #caminho-apis-c-completo não encontrado", file=sys.stderr)
        return []

    # Remove elementos não sancionáveis
    for unwanted_class in ["dec-nota-redacao", "resumo-executivo", "caminho-completo__header",
                            "caminho-completo__aviso-orientacoes"]:
        for el in block.find_all(class_=unwanted_class):
            el.decompose()

    items = []
    for el in block.find_all(True, recursive=True):
        classes = el.get("class", [])

        if "dec-nota-redacao" in classes:
            continue
        if "dec-section__subtitle" in classes:
            continue
        if "dec-section__num" in classes:
            continue

        if "dec-norma__ementa" in classes:
            txt = clean_text(el.get_text())
            if txt:
                tipo = "ementa_italico" if "dec-norma__ementa--italico" in classes else "ementa"
                items.append((tipo, txt))
            continue

        if "dec-norma__autor" in classes:
            txt = clean_text(el.get_text())
            if txt:
                items.append(("autor", txt))
            continue

        if "dec-norma__bloco-title" in classes:
            txt = clean_text(el.get_text())
            if txt:
                items.append(("bloco_title", txt))
            continue

        if "dec-section__title" in classes:
            txt = clean_text(el.get_text())
            if txt:
                items.append(("section_title", txt))
            continue

        if "dec-artigo" in classes:
            txt = clean_text(el.get_text())
            if txt:
                items.append(("artigo", txt))
            continue

        if "dec-paragrafo" in classes:
            txt = clean_text(el.get_text())
            if txt:
                items.append(("paragrafo", txt))
            continue

        if "dec-considerandos" in classes:
            for li in el.find_all("li", recursive=False):
                txt = clean_text(li.get_text())
                if txt and not txt.startswith("(sem considerando"):
                    items.append(("considerando", txt))
            continue

        if "dec-incisos" in classes:
            for li in el.find_all("li"):
                txt = clean_text(li.get_text())
                if txt:
                    items.append(("inciso", txt))
            continue

    return items


# ── Montar DOCX ────────────────────────────────────────────────────────────────
doc = gw.create_base_doc()

# Capa
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run("PREFEITURA MUNICIPAL DE JOINVILLE")
run.bold = True
run.font.size = Pt(14)
gw.set_paragraph_format(p_title, space_before=0, space_after=6)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run("MINUTA DE DECRETO")
r2.font.size = Pt(12)
gw.set_paragraph_format(p_sub, space_before=0, space_after=4)

p_nota = doc.add_paragraph()
p_nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_nota.add_run("Versão consolidada para análise jurídica")
r3.italic = True
r3.font.size = Pt(10)
gw.set_paragraph_format(p_nota, space_before=0, space_after=4)

p_org = doc.add_paragraph()
p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p_org.add_run("BRZ Capacitação × Consultoria SEBRAE/SC · maio/2026")
r4.font.size = Pt(10)
gw.set_paragraph_format(p_org, space_before=0, space_after=20)

# Título
p_titulo_dec = doc.add_paragraph()
p_titulo_dec.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = p_titulo_dec.add_run("DECRETO DOS ARRANJOS PROMOTORES DE INOVAÇÃO (APIs)")
rt.bold = True
rt.font.size = Pt(12)
gw.set_paragraph_format(p_titulo_dec, space_before=20, space_after=18)

apis_c_items = extract_apis_c_sections()
print(f"APIs (Caminho C): {len(apis_c_items)} elementos extraídos")
gw.render_items(doc, apis_c_items)

out_path = OUT_DIR / "Joinville - R-C - Decreto APIs.docx"
doc.save(str(out_path))
print(f"\n✓ Salvo: {out_path}")
